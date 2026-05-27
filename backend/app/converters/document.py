import json
import io
import csv
from pathlib import Path


def convert_document(input_path: Path, source_ext: str, target_ext: str) -> Path:
    output_path = input_path.with_suffix(f".{target_ext}")
    s, t = source_ext, target_ext

    if s == "pdf":
        return _convert_from_pdf(input_path, t, output_path)
    if s == "docx":
        return _convert_from_docx(input_path, t, output_path)
    if s == "xlsx":
        return _convert_from_xlsx(input_path, t, output_path)
    if s == "html":
        return _convert_from_html(input_path, t, output_path)
    if s == "md":
        return _convert_from_md(input_path, t, output_path)
    if s == "txt" and t in ("docx", "pdf"):
        return _txt_convert(input_path, t, output_path)

    raise ValueError(f"Unsupported conversion: {s} -> {t}")


def _convert_from_pdf(input_path: Path, target: str, output_path: Path) -> Path:
    import pdfplumber

    if target == "txt":
        text = ""
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        output_path.write_text(text, encoding="utf-8")
        return output_path

    if target == "html":
        parts = ['<!DOCTYPE html><html><head><meta charset="utf-8"><title>PDF</title></head><body>']
        with pdfplumber.open(input_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    parts.append(f'<h2>Page {i+1}</h2>')
                    for line in text.split("\n"):
                        parts.append(f'<p>{line}</p>')
        parts.append('</body></html>')
        output_path.write_text("\n".join(parts), encoding="utf-8")
        return output_path

    if target == "docx":
        from docx import Document
        doc = Document()
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    for line in text.split("\n"):
                        if line.strip():
                            doc.add_paragraph(line.strip())
        doc.save(str(output_path))
        return output_path

    if target == "md":
        text = ""
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n\n"
        output_path.write_text(text, encoding="utf-8")
        return output_path

    if target in ("csv", "xlsx", "json"):
        return _pdf_extract_table(input_path, target, output_path)

    raise ValueError(f"PDF -> {target} not supported")


def _pdf_extract_table(input_path: Path, target: str, output_path: Path) -> Path:
    import pdfplumber

    all_tables = []
    with pdfplumber.open(input_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                if table and len(table) > 1:
                    headers = table[0]
                    for row in table[1:]:
                        if row and any(cell for cell in row):
                            all_tables.append(dict(zip(headers, row)))

    if target == "csv":
        if all_tables:
            buf = io.StringIO()
            writer = csv.DictWriter(buf, fieldnames=all_tables[0].keys())
            writer.writeheader()
            writer.writerows(all_tables)
            output_path.write_text(buf.getvalue(), encoding="utf-8")
        else:
            output_path.write_text("", encoding="utf-8")
        return output_path

    if target == "json":
        output_path.write_text(json.dumps(all_tables, ensure_ascii=False, indent=2), encoding="utf-8")
        return output_path

    if target == "xlsx":
        from .data import csv_to_xlsx
        if all_tables:
            import pandas as pd
            df = pd.DataFrame(all_tables)
            df.to_excel(output_path, index=False)
        else:
            import pandas as pd
            pd.DataFrame().to_excel(output_path, index=False)
        return output_path

    raise ValueError(f"PDF table -> {target}")


def _convert_from_docx(input_path: Path, target: str, output_path: Path) -> Path:
    from docx import Document
    doc = Document(input_path)

    if target == "txt":
        text = "\n".join(p.text for p in doc.paragraphs)
        output_path.write_text(text, encoding="utf-8")
        return output_path

    if target == "md":
        md_lines = []
        for p in doc.paragraphs:
            style = p.style.name.lower() if p.style else ""
            text = p.text
            if not text:
                md_lines.append("")
            elif "heading 1" in style:
                md_lines.append(f"# {text}")
            elif "heading 2" in style:
                md_lines.append(f"## {text}")
            elif "heading 3" in style:
                md_lines.append(f"### {text}")
            else:
                md_lines.append(text)
        output_path.write_text("\n\n".join(md_lines), encoding="utf-8")
        return output_path

    if target == "html":
        html_parts = ['<!DOCTYPE html><html><head><meta charset="utf-8"></head><body>']
        for p in doc.paragraphs:
            style = p.style.name.lower() if p.style else ""
            text = p.text
            if not text:
                html_parts.append("<br>")
            elif "heading 1" in style:
                html_parts.append(f"<h1>{text}</h1>")
            elif "heading 2" in style:
                html_parts.append(f"<h2>{text}</h2>")
            elif "heading 3" in style:
                html_parts.append(f"<h3>{text}</h3>")
            else:
                html_parts.append(f"<p>{text}</p>")
        html_parts.append("</body></html>")
        output_path.write_text("\n".join(html_parts), encoding="utf-8")
        return output_path

    if target == "pdf":
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)
        for p in doc.paragraphs:
            text = p.text
            if text.strip():
                pdf.multi_cell(0, 6, text)
            else:
                pdf.ln(6)
        pdf.output(str(output_path))
        return output_path

    if target == "xml":
        xml_parts = ['<?xml version="1.0" encoding="UTF-8"?>', '<document>']
        for para in doc.paragraphs:
            xml_parts.append(f'  <paragraph>{para.text}</paragraph>')
        xml_parts.append('</document>')
        output_path.write_text("\n".join(xml_parts), encoding="utf-8")
        return output_path

    raise ValueError(f"DOCX -> {target} not supported")


def _convert_from_xlsx(input_path: Path, target: str, output_path: Path) -> Path:
    import pandas as pd
    df = pd.read_excel(input_path)

    if target == "csv":
        df.to_csv(output_path, index=False)
    elif target == "json":
        df.to_json(output_path, orient="records", force_ascii=False, indent=2)
    elif target == "sql":
        table_name = input_path.stem.replace(" ", "_").replace("-", "_")
        lines = []
        cols = ", ".join(df.columns)
        for _, row in df.iterrows():
            vals = ", ".join(f"'{str(v).replace(chr(39), chr(39)+chr(39))}'" for v in row.values)
            lines.append(f"INSERT INTO {table_name} ({cols}) VALUES ({vals});")
        output_path.write_text("\n".join(lines), encoding="utf-8")
    elif target == "pdf":
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=10)
        for _, row in df.iterrows():
            line = " | ".join(str(v) for v in row.values)
            pdf.cell(0, 7, line, ln=True)
        pdf.output(str(output_path))
    elif target == "txt":
        df.to_csv(output_path, index=False, sep="\t")
    else:
        raise ValueError(f"XLSX -> {target}")
    return output_path


def _convert_from_html(input_path: Path, target: str, output_path: Path) -> Path:
    content = input_path.read_text(encoding="utf-8")

    if target == "txt":
        from html.parser import HTMLParser
        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text = []
            def handle_data(self, data):
                self.text.append(data)
        extractor = TextExtractor()
        extractor.feed(content)
        output_path.write_text(" ".join(extractor.text), encoding="utf-8")
        return output_path

    if target == "docx":
        from docx import Document
        from html.parser import HTMLParser
        doc = Document()
        current_text = []

        class HTMLHandler(HTMLParser):
            def handle_data(self, data):
                if data.strip():
                    current_text.append(data.strip())
            def handle_endtag(self, tag):
                if tag in ('p', 'h1', 'h2', 'h3', 'h4', 'div', 'br', 'li'):
                    if current_text:
                        doc.add_paragraph(" ".join(current_text))
                        current_text.clear()

        handler = HTMLHandler()
        handler.feed(content)
        if current_text:
            doc.add_paragraph(" ".join(current_text))
        doc.save(str(output_path))
        return output_path

    if target == "md":
        import re
        text = re.sub(r'<[^>]+>', '', content)
        output_path.write_text(text.strip(), encoding="utf-8")
        return output_path

    raise ValueError(f"HTML -> {target}")


def _convert_from_md(input_path: Path, target: str, output_path: Path) -> Path:
    import markdown
    content = input_path.read_text(encoding="utf-8")

    if target == "html":
        html = markdown.markdown(content)
        full = f'<!DOCTYPE html><html><head><meta charset="utf-8"></head><body>\n{html}\n</body></html>'
        output_path.write_text(full, encoding="utf-8")
        return output_path

    if target == "docx":
        from docx import Document
        html = markdown.markdown(content)
        from html.parser import HTMLParser
        doc = Document()
        current = []

        class H(HTMLParser):
            def handle_data(self, d):
                if d.strip(): current.append(d.strip())
            def handle_endtag(self, t):
                if t in ('p','h1','h2','h3','h4','li') and current:
                    doc.add_paragraph(" ".join(current))
                    current.clear()

        H().feed(html)
        if current:
            doc.add_paragraph(" ".join(current))
        doc.save(str(output_path))
        return output_path

    if target == "txt":
        from html.parser import HTMLParser
        html = markdown.markdown(content)
        class T(HTMLParser):
            def __init__(self):
                super().__init__()
                self.t = []
            def handle_data(self, d):
                self.t.append(d)
        ex = T()
        ex.feed(html)
        output_path.write_text("\n".join(ex.t), encoding="utf-8")
        return output_path

    if target == "pdf":
        from fpdf import FPDF
        html = markdown.markdown(content)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)
        for line in html.split("\n"):
            clean = line.replace("<p>","").replace("</p>","").replace("<h1>","").replace("</h1>","").replace("<h2>","").replace("</h2>","").replace("<h3>","").replace("</h3>","")
            if clean.strip():
                pdf.multi_cell(0, 6, clean.strip())
        pdf.output(str(output_path))
        return output_path

    raise ValueError(f"MD -> {target}")


def _txt_convert(input_path: Path, target: str, output_path: Path) -> Path:
    if target == "docx":
        from docx import Document
        doc = Document()
        for line in input_path.read_text(encoding="utf-8").split("\n"):
            doc.add_paragraph(line)
        doc.save(str(output_path))
        return output_path

    if target == "pdf":
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)
        for line in input_path.read_text(encoding="utf-8").split("\n"):
            pdf.cell(0, 7, line, ln=True)
        pdf.output(str(output_path))
        return output_path
